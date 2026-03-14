"""
File-based ingester: reads text from TXT, PDF, and DOCX files.
"""
import uuid
from pathlib import Path

from app.ingestion.base import BaseIngester, IngestedDocument
from app.logger import get_logger
from app.utils.file_utils import detect_mime_type, file_size_kb, is_supported_file

logger = get_logger(__name__)


class FileIngester(BaseIngester):
    """Ingests documents from the local filesystem."""

    async def can_handle(self, source: str) -> bool:
        p = Path(source)
        return p.exists() and p.is_file() and is_supported_file(p)

    async def ingest(self, source: str, metadata: dict | None = None) -> IngestedDocument:
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {source}")

        mime = detect_mime_type(path)
        size = file_size_kb(path)
        logger.info("Ingesting file: %s (%s, %.1f KB)", path.name, mime, size)

        raw_text = await self._extract_text(path, mime)

        return IngestedDocument(
            id=str(uuid.uuid4()),
            source_type="file",
            raw_text=raw_text,
            filename=path.name,
            mime_type=mime,
            size_kb=size,
            metadata=metadata or {},
        )

    async def _extract_text(self, path: Path, mime: str) -> str:
        """Dispatch to the appropriate text extractor based on MIME type."""
        if mime == "text/plain" or mime == "text/markdown":
            return self._read_text_file(path)
        elif mime == "application/pdf":
            return self._read_pdf(path)
        elif mime in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            return self._read_docx(path)
        else:
            raise ValueError(f"Unsupported MIME type for text extraction: {mime}")

    def _read_text_file(self, path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="replace")

    def _read_pdf(self, path: Path) -> str:
        """Extract text from PDF. Requires pypdf (optional dependency)."""
        try:
            import pypdf  # type: ignore
            reader = pypdf.PdfReader(str(path))
            return "\n\n".join(
                page.extract_text() or "" for page in reader.pages
            )
        except ImportError:
            raise ImportError(
                "pypdf is required for PDF ingestion. Install with: pip install pypdf"
            )

    def _read_docx(self, path: Path) -> str:
        """Extract text from DOCX. Requires python-docx (optional dependency)."""
        try:
            import docx  # type: ignore
            doc = docx.Document(str(path))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX ingestion. Install with: pip install python-docx"
            )
